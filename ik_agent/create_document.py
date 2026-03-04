"""
Yorglass IK Prosedur Dokumani Olusturucu.

python-docx ile programatik olarak Yorglass IK prosedur dokumanini olusturur.
Bu script bir kez calistirilir ve ik_agent/data/ altina .docx dosyasi uretir.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_ik_document():
    """Yorglass IK Prosedur Dokumanini olusturur."""
    doc = Document()

    # Stil ayarlari
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # ── KAPAK ──
    title = doc.add_heading("YORGLASS CAM SANAYI A.S.", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading("Insan Kaynaklari Prosedur ve Politikalari El Kitabi", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Dokuman No: YRG-IK-POL-001\n").bold = True
    meta.add_run("Revizyon: 3.0\n")
    meta.add_run("Yururluk Tarihi: 01.01.2024\n")
    meta.add_run("Hazirlayan: Insan Kaynaklari Departmani\n")
    meta.add_run("Onaylayan: Genel Mudur")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # BOLUM 1: GENEL BILGILER
    # ══════════════════════════════════════════
    doc.add_heading("1. Genel Bilgiler", level=1)

    doc.add_heading("1.1 Amac", level=2)
    doc.add_paragraph(
        "Bu dokuman, Yorglass Cam Sanayi A.S. bunyesinde calisan tum personelin "
        "tabi oldugu insan kaynaklari politikalarini, prosedurlerini ve calisma "
        "kurallarini belirlemek amaciyla hazirlanmistir. Dokumandaki hukumler "
        "4857 sayili Is Kanunu ve ilgili mevzuat cercevesinde duzenlenmistir."
    )

    doc.add_heading("1.2 Kapsam", level=2)
    doc.add_paragraph(
        "Bu el kitabi, Yorglass Cam Sanayi A.S.'nin tum lokasyonlarinda "
        "(merkez ofis, uretim tesisleri ve saha operasyonlari) gorev yapan "
        "tam zamanli, yari zamanli ve sozlesmeli tum calisanlari kapsar. "
        "Stajyerler ve gecici isci statusundeki personel icin ayrica belirtilen "
        "ozel hukumler saklıdır."
    )

    doc.add_heading("1.3 Sirket Hakkinda", level=2)
    doc.add_paragraph(
        "Yorglass Cam Sanayi A.S., 1982 yilinda kurulan, Turkiye'nin oncu "
        "cam isleme ve uretim sirketlerinden biridir. Sirketimiz duz cam, "
        "temperli cam, lamine cam, ayna ve yalitim cami uretimi konusunda "
        "faaliyet gostermektedir. Merkezimiz Istanbul'da olup, Gebze ve "
        "Bolu'da uretim tesislerimiz bulunmaktadir. Yaklasik 450 calisan "
        "ile sektorde hizmet vermekteyiz."
    )

    # ══════════════════════════════════════════
    # BOLUM 2: ISE GIRIS SURECI
    # ══════════════════════════════════════════
    doc.add_heading("2. Ise Giris Sureci", level=1)

    doc.add_heading("2.1 Basvuru ve Mulakat", level=2)
    doc.add_paragraph(
        "Yorglass'a is basvurusu, sirket web sitesi, kariyer portallari veya "
        "dogrudan IK departmanina yapilabilir. Basvurular IK birimi tarafindan "
        "degerlendirilir ve uygun bulunan adaylar mulakat surecine davet edilir. "
        "Mulakat sureci en az iki asamadan olusur: IK on gorusmesi ve teknik/"
        "yonetici mulakati. Bazi pozisyonlar icin yetenek testi veya uygulama "
        "sinavi yapilabilir."
    )

    doc.add_heading("2.2 Gerekli Belgeler", level=2)
    doc.add_paragraph(
        "Ise giris icin asagidaki belgelerin IK departmanina teslim edilmesi "
        "zorunludur: Nufus cuzdani fotokopisi (2 adet), ikametgah belgesi, "
        "adli sicil kaydi, saglik raporu (is yeri hekiminden onaylanmis), "
        "diploma fotokopisi, 4 adet vesikalik fotograf, SGK ise giris bildirgesi, "
        "banka hesap bilgileri ve varsa is deneyimi belgeleri. Eksik belge ile "
        "ise baslama mumkun degildir."
    )

    doc.add_heading("2.3 Deneme Suresi", level=2)
    doc.add_paragraph(
        "Tum yeni calisanlar icin 2 (iki) aylik deneme suresi uygulanir. "
        "4857 sayili Is Kanunu'nun 15. maddesi geregince, deneme suresi "
        "icinde taraflardan her biri is sozlesmesini bildirim suresi olmaksizin "
        "ve tazminatsiz feshedebilir. Deneme suresi sonunda calisan ve yoneticisi "
        "tarafindan performans degerlendirmesi yapilir. Basarili bulunan calisanlar "
        "kadroya alinir."
    )

    doc.add_heading("2.4 Oryantasyon", level=2)
    doc.add_paragraph(
        "Ise yeni baslayan her calisan icin 3 gunluk oryantasyon programi "
        "uygulanir. Bu program sirket tanitimi, is guvenligi egitimi, "
        "departman tanitimi, IT sistemleri egitimi ve IK prosedurlerinin "
        "anlatilmasini kapsar. Uretim tesislerinde ise baslayacak personel "
        "icin oryantasyon suresi 5 is gunune kadar uzatilabilir."
    )

    # ══════════════════════════════════════════
    # BOLUM 3: CALISMA SAATLERI
    # ══════════════════════════════════════════
    doc.add_heading("3. Calisma Saatleri", level=1)

    doc.add_heading("3.1 Haftalik Calisma Suresi", level=2)
    doc.add_paragraph(
        "Yorglass'ta haftalik calisma suresi 45 saattir. Bu sure, haftanin "
        "5 is gunune esit olarak dagitilir. Ofis personeli icin mesai saatleri "
        "08:30 - 17:30 arasindadir. Uretim tesislerinde vardiyali calisma sistemi "
        "uygulanmakta olup, vardiya saatleri 06:00-14:00, 14:00-22:00 ve "
        "22:00-06:00 olarak belirlenmistir."
    )

    doc.add_heading("3.2 Mola Sureleri", level=2)
    doc.add_paragraph(
        "Gunluk calisma suresinin ortasinda 1 saat oglen yemegi molasi verilir. "
        "Ofis personeli icin oglen molasi 12:00 - 13:00 arasindadir. Uretim "
        "tesislerinde vardiya programina gore mola sureleri ayarlanir. Buna ek "
        "olarak sabah ve ogle sonrasi 15'er dakikalik cay/kahve molasi verilir. "
        "Mola sureleri calisma suresinden sayilmaz."
    )

    doc.add_heading("3.3 Devam ve Yoklama", level=2)
    doc.add_paragraph(
        "Tum calisanlar ise giris ve cikis saatlerini kart basma sistemi "
        "veya dijital yoklama uygulamasi ile kayit altina alir. 3 defadan "
        "fazla uyarisiz gec gelme durumunda yazili uyari yapilir. Mazeretsiz "
        "ise gelmeme 2 ardisik is gununu astigi takdirde disiplin sureci "
        "baslatilir. Aylik devam raporlari IK departmani tarafindan takip edilir."
    )

    # ══════════════════════════════════════════
    # BOLUM 4: YILLIK IZIN
    # ══════════════════════════════════════════
    doc.add_heading("4. Yillik Izin", level=1)

    doc.add_heading("4.1 Izin Hakki ve Suresi", level=2)
    doc.add_paragraph(
        "4857 sayili Is Kanunu'nun 53. maddesi geregince, yillik ucretli izin "
        "hakki asagidaki sekilde belirlenmistir: 1 yildan 5 yila kadar (5 yil "
        "dahil) kidem suresi olan calisanlar yilda 14 gun, 5 yildan fazla 15 "
        "yildan az kidem suresi olan calisanlar yilda 20 gun, 15 yil (dahil) ve "
        "uzerinde kidem suresi olan calisanlar yilda 26 gun yillik izin hakkina "
        "sahiptir. 18 yasindan kucuk ve 50 yasindan buyuk calisanlar icin yillik "
        "izin suresi 20 is gununden az olamaz."
    )

    doc.add_heading("4.2 Izin Kullanim Kurallari", level=2)
    doc.add_paragraph(
        "Yillik izin talepleri, izin tarihinden en az 5 is gunu once IK "
        "sistemine girilmeli ve birim yoneticisi tarafindan onaylanmalidir. "
        "Yillik izin en fazla 3 parcaya bolunerek kullanilabilir. Ancak her "
        "bir izin parcasi 5 is gunundan az olamaz. Toplu izin donemlerinde "
        "(Agustos, yilbasi) departman bazinda koordinasyon yapilmasi zorunludur. "
        "Kullanilmayan yillik izinler bir sonraki yila devreder, ancak 2 yildan "
        "fazla birikmis izin kullanilmasi IK departmani tarafindan planlanir."
    )

    doc.add_heading("4.3 Izin Hesaplama ve Takip", level=2)
    doc.add_paragraph(
        "Yillik izin hakki, calisanin ise baslama tarihinden itibaren 1 yillik "
        "kidem sureleri tamamlandikca hak edilir. Izin hesaplamalari IK otomasyon "
        "sistemi uzerinden otomatik yapilir. Her calisanin kalan izin gunu "
        "bilgisi IK portalindan sorgulanabilir. Izin donus tarihi itibariyle "
        "calisan ise baslamak zorundadir; mazeret bildirmeksizin donmeyen calisan "
        "hakkinda disiplin islemi uygulanir."
    )

    # ══════════════════════════════════════════
    # BOLUM 5: SAGLIK RAPORU
    # ══════════════════════════════════════════
    doc.add_heading("5. Saglik Raporu", level=1)

    doc.add_heading("5.1 Rapor Suresi ve Bildirim", level=2)
    doc.add_paragraph(
        "Calisan hastalandiginda ayni gun icinde (mesai baslamadan once veya en gec "
        "mesai basladigindan itibaren 2 saat icinde) birim yoneticisini ve IK "
        "departmanini bilgilendirmekle yukumludur. Bildirim telefon veya mesaj "
        "yoluyla yapilabilir. Saglik raporu 3 is gununun altinda ise rapor bitiminde "
        "isleme tabi tutulur. 3 is gunu ve uzeri raporlarda SGK'ya bildirim yapilir "
        "ve calisanin rapor suresi boyunca ucreti SGK tarafindan odenir."
    )

    doc.add_heading("5.2 Belgeleme", level=2)
    doc.add_paragraph(
        "Tum saglik raporlari, rapor bitiminde ise donus yapildiginda IK "
        "departmanina teslim edilmelidir. Rapor; devlet hastanesi, universite "
        "hastanesi veya SGK anlasmalı ozel saglik kuruluslarinan alinmis olmalidir. "
        "Ayni hastalik nedeniyle yil icerisinde 3 defadan fazla kisa sureli rapor "
        "alan calisanlar icin IK departmani tarafindan saglik durumu degerlendirmesi "
        "yapilabilir. Raporlu olunan surelerde calisan baska bir iste calismamakla "
        "yukumludur."
    )

    doc.add_heading("5.3 Is Kazasi Raporu", level=2)
    doc.add_paragraph(
        "Is kazasi meydana gelmesi halinde, kaza aninda veya en gec 3 gun icinde "
        "Is Guvenligi Birimine ve IK departmanina bildirim yapilmasi zorunludur. "
        "Is kazasi raporu ayrica SGK'ya 3 is gunu icinde bildirilir. Is kazasi "
        "nedeniyle alinan raporlarin tamami SGK tarafindan karsilanir. Calisan "
        "is kazasi sonrasi ise donuste is yeri hekiminden uygunluk onay almalidir."
    )

    # ══════════════════════════════════════════
    # BOLUM 6: MAZERET IZNI
    # ══════════════════════════════════════════
    doc.add_heading("6. Mazeret Izni", level=1)

    doc.add_heading("6.1 Ucretli Mazeret Izinleri", level=2)
    doc.add_paragraph(
        "4857 sayili Is Kanunu ve toplu is sozlesmesi kapsaminda calisanlara "
        "asagidaki ucretli mazeret izinleri verilir: Evlilik izni 3 is gunu, "
        "cocuk dunyaya gelmesi (baba icin) 5 is gunu, anne veya baba, es, "
        "kardes veya cocugun vefati halinde 3 is gunu, calisanin evlat "
        "edinmesi halinde 3 is gunu. Bu izinler olay tarihinden itibaren "
        "kullanilmalidir ve sonraya ertelenemez."
    )

    doc.add_heading("6.2 Diger Mazeret Izinleri", level=2)
    doc.add_paragraph(
        "Tasinma izni olarak yilda 1 gun ucretli izin verilir. Askerlik "
        "yoklamasi veya celp icin gerekli surece ucretli izin verilir. "
        "Mahkeme celbi, taniklik veya resmi kurum islemleri icin belgelenmek "
        "kaydiyla ucretli izin verilir. Yilda toplam 5 is gununu asmamak "
        "uzere, IK departmaninin uygun gormesi halinde ozel mazeret izni "
        "verilebilir. 5 gunu asan mazeret izinleri ucretsiz izin statusunde "
        "degerlendirilir."
    )

    # ══════════════════════════════════════════
    # BOLUM 7: UCRETSIZ IZIN
    # ══════════════════════════════════════════
    doc.add_heading("7. Ucretsiz Izin", level=1)

    doc.add_heading("7.1 Ucretsiz Izin Kosullari", level=2)
    doc.add_paragraph(
        "Calisanlar, yillik ucretli izin haklarini tukettikten sonra, "
        "yazili olarak ucretsiz izin talebinde bulunabilir. Ucretsiz izin "
        "talebi, en az 15 gun onceden IK departmanina iletilmeli ve birim "
        "yoneticisi ile IK Muduru tarafindan onaylanmalidir. Ucretsiz izin "
        "suresi bir takvim yilinda toplam 30 is gununu gecemez. Is akisinin "
        "aksamamasi icin ucretsiz izin donemlerinde yerine gecici personel "
        "gorevlendirilebilir."
    )

    doc.add_heading("7.2 Dogum Sonrasi Ucretsiz Izin", level=2)
    doc.add_paragraph(
        "4857 sayili Is Kanunu'nun 74. maddesi geregince, kadin calisanlar "
        "dogum sonrasi ucretli izin suresinin bitiminden itibaren 6 aya kadar "
        "ucretsiz izin talep edebilir. Bu izin suresinde calisanin is sozlesmesi "
        "askiya alinir ve kidem suresi islemeye devam eder. Ucretsiz izin "
        "suresince calisanin SGK primleri odenmez. Izin bitiminde calisan ayni "
        "veya esdeger pozisyona iade edilir."
    )

    # ══════════════════════════════════════════
    # BOLUM 8: FAZLA MESAI
    # ══════════════════════════════════════════
    doc.add_heading("8. Fazla Mesai", level=1)

    doc.add_heading("8.1 Fazla Mesai Tanimlari", level=2)
    doc.add_paragraph(
        "Haftalik 45 saati asan calisma sureleri fazla calisma (fazla mesai) "
        "olarak kabul edilir. 4857 sayili Is Kanunu'nun 41. maddesi geregince, "
        "fazla calisma ucreti normal calisma ucretinin %50 fazlasiyla odenir. "
        "Tatil gunlerinde yapilan calismalarda ucret %100 zamlı olarak hesaplanir. "
        "Yillik fazla calisma suresi 270 saati gecemez."
    )

    doc.add_heading("8.2 Onay Sureci", level=2)
    doc.add_paragraph(
        "Fazla mesai yapilabilmesi icin birim yoneticisinin onceden yazili "
        "onay vermesi gereklidir. Fazla mesai formu, mesai yapilacak tarihten "
        "en az 1 is gunu once doldurulup imzalanmali ve IK departmanina "
        "iletilmelidir. Acil durumlar haricinde onceden onay alinmamis fazla "
        "mesai kabul edilmez. Aylik fazla mesai sureleri IK departmani tarafindan "
        "takip edilir ve bordro birimine raporlanir."
    )

    doc.add_heading("8.3 Fazla Mesai Hesaplama", level=2)
    doc.add_paragraph(
        "Fazla mesai ucreti su formule gore hesaplanir: Saatlik ucret = "
        "Aylik brut ucret / 225. Fazla mesai ucreti = Saatlik ucret x 1.5 x "
        "fazla calisilan saat sayisi. Hafta tatili ve resmi tatillerde yapilan "
        "calismalarda: Saatlik ucret x 2.0 x calisilan saat sayisi. Fazla mesai "
        "yerine serbest zaman kullanilmasi istenirse, her 1 saat fazla calisma "
        "icin 1.5 saat serbest zaman verilir."
    )

    # ══════════════════════════════════════════
    # BOLUM 9: KIYAFET KURALLARI
    # ══════════════════════════════════════════
    doc.add_heading("9. Kiyafet Kurallari", level=1)

    doc.add_heading("9.1 Ofis Personeli", level=2)
    doc.add_paragraph(
        "Ofis calisanlari is ortamina uygun, temiz ve bakimli kiyafetler "
        "giymelidir. Erkek calisanlar icin gomlek, kumaş pantolon ve kapalı "
        "ayakkabi tercih edilmelidir. Kadin calisanlar icin is ortamina uygun "
        "sade kiyafetler kabul edilir. Cuma gunleri serbest kiyafet (smart casual) "
        "uygulamasi gecerlidir. Musteri ziyareti veya toplantı gunlerinde resmi "
        "kiyafet zorunludur."
    )

    doc.add_heading("9.2 Uretim Tesisi Personeli", level=2)
    doc.add_paragraph(
        "Uretim tesislerinde calisan tum personel, sirket tarafindan verilen "
        "is elbisesi, celik burunlu is ayakkabisi, baret, koruyucu gozluk ve "
        "eldiven gibi kisisel koruyucu donanimlari (KKD) kullanmak zorundadir. "
        "KKD'siz uretim alanina girmek kesinlikle yasaktir. KKD ihlali ilk "
        "seferinde yazili uyari, tekrarinda disiplin cezasi ile sonuclanir. "
        "Yipranmis veya hasarli KKD'ler derhal IK veya Is Guvenligi birimine "
        "bildirilerek yenisiyle degistirilmelidir."
    )

    # ══════════════════════════════════════════
    # BOLUM 10: DISIPLIN PROSEDURU
    # ══════════════════════════════════════════
    doc.add_heading("10. Disiplin Proseduru", level=1)

    doc.add_heading("10.1 Disiplin Adimlari", level=2)
    doc.add_paragraph(
        "Yorglass'ta disiplin proseduru kademeli olarak uygulanir: "
        "1. Adim: Sozlu uyari - Yonetici tarafindan calisan ile bire bir "
        "gorusme yapilir ve durum kayit altina alinir. "
        "2. Adim: Yazili uyari - IK departmani tarafindan resmi yazili uyari "
        "verilir ve calisanin ozluk dosyasina islenir. "
        "3. Adim: Son uyari - Calisana disiplin kurulu karari ile son uyari "
        "verilir ve fesih ihtimali bildirilir. "
        "4. Adim: Is akdinin feshi - Isverenin hakli fesih nedenleri "
        "(Is Kanunu madde 25/II) cercevesinde is sozlesmesi sonlandirilir."
    )

    doc.add_heading("10.2 Agir Disiplin Suclari", level=2)
    doc.add_paragraph(
        "Asagidaki durumlarda kademeli disiplin proseduru uygulanmaksizin "
        "dogrudan is akdi feshi yapilabilir: Hirsizlik, zimmet veya dolandiricilik, "
        "is yerinde siddet veya tehdit, alkollu veya uyusturucu madde etkisinde "
        "ise gelme, kasitli olarak is ekipmanina zarar verme, is sirlarina ihanet "
        "veya gizlilik ihali, cinsel taciz veya ayrimcilik, 3 ardisik is gunu "
        "mazeretsiz devamsizlik. Bu durumlarda 4857 sayili Is Kanunu'nun 25/II "
        "maddesi kapsaminda hakli fesih uygulanir."
    )

    # ══════════════════════════════════════════
    # BOLUM 11: PERFORMANS DEGERLENDIRME
    # ══════════════════════════════════════════
    doc.add_heading("11. Performans Degerlendirme", level=1)

    doc.add_heading("11.1 Degerlendirme Donemleri", level=2)
    doc.add_paragraph(
        "Yorglass'ta performans degerlendirmesi yilda 2 kez yapilir: "
        "Ara degerlendirme Haziran ayinda, yil sonu degerlendirmesi Aralik "
        "ayinda gerceklestirilir. Degerlendirme, calisan ile yonetici arasinda "
        "yuz yuze gorusme seklinde yapilir. Hem calisanin oz degerlendirmesi "
        "hem de yoneticinin degerlendirmesi dikkate alinir."
    )

    doc.add_heading("11.2 Degerlendirme Kriterleri", level=2)
    doc.add_paragraph(
        "Performans degerlendirmesinde su kriterler esas alinir: "
        "Is kalitesi ve verimlilik (%30), hedef gerceklestirme orani (%25), "
        "takim calismasina katki (%15), inisiyatif ve problem cozme becerisi (%15), "
        "mesleki gelisim ve ogrenme istegi (%10), devam ve dakiklik (%5). "
        "Her kriter 1-5 arasi puanlanir. Genel puan ortalamasi 4.0 ve uzeri "
        "'Ustun Basari', 3.0-3.9 arasi 'Basarili', 2.0-2.9 arasi 'Gelisime Acik', "
        "2.0 alti 'Yetersiz' olarak degerlendirilir."
    )

    doc.add_heading("11.3 Terfi ve Ucret Artisi", level=2)
    doc.add_paragraph(
        "Yil sonu performans degerlendirmesinde 'Ustun Basari' veya 'Basarili' "
        "alan calisanlar terfi ve ucret artisi icin degerlendirilir. Terfi karari "
        "departman yoneticisi, IK departmani ve ust yonetimin ortak karari ile "
        "verilir. Ucret artislari, performans puani, enflasyon orani ve piyasa "
        "koşulları dikkate alinarak belirlenir. 'Yetersiz' puan alan calisanlar "
        "icin 3 aylik performans iyilestirme plani uygulanir."
    )

    # ══════════════════════════════════════════
    # BOLUM 12: EGITIM VE GELISIM
    # ══════════════════════════════════════════
    doc.add_heading("12. Egitim ve Gelisim", level=1)

    doc.add_heading("12.1 Egitim Haklari", level=2)
    doc.add_paragraph(
        "Yorglass, calisanlarinin mesleki gelisimine onem verir. Her calisan "
        "yilda en az 20 saat mesleki egitim alma hakkina sahiptir. Egitim "
        "programlari IK departmani tarafindan yillik egitim plani cercevesinde "
        "organize edilir. Zorunlu egitimler (is guvenligi, kalite yonetimi, "
        "cevre bilinci) yillik egitim saatine dahildir."
    )

    doc.add_heading("12.2 Dis Egitim ve Sertifikasyon", level=2)
    doc.add_paragraph(
        "Calisanlar, isi ile ilgili dis egitim ve sertifikasyon programlarina "
        "katilmak icin IK departmanina yazili basvuruda bulunabilir. Sirket "
        "tarafindan onaylanan dis egitimlerin masrafi (kurs ucreti, ulasim, "
        "konaklama) sirket tarafindan karsilanir. Dis egitim suresi mesai "
        "saatinden sayilir. Sertifikasyon sinavlarinda basarili olan calisanlara "
        "basari primi verilir. Egitim masrafi 5.000 TL'yi asan durumlarda "
        "calisan, egitim taahhutnamesi imzalar (2 yil calisma taahhüdü)."
    )

    # ══════════════════════════════════════════
    # BOLUM 13: UZAKTAN CALISMA
    # ══════════════════════════════════════════
    doc.add_heading("13. Uzaktan Calisma", level=1)

    doc.add_heading("13.1 Uygunluk Kosullari", level=2)
    doc.add_paragraph(
        "Uzaktan calisma, is niteligi uygun olan ofis pozisyonlarinda, "
        "birim yoneticisi ve IK departmaninin onayiyla uygulanabilir. "
        "Uretim, depo, lojistik ve saha operasyonlari gibi fiziksel varlik "
        "gerektiren pozisyonlar uzaktan calisma kapsamina girmez. Uzaktan "
        "calisma haftada en fazla 2 gun olarak uygulanabilir. Calisan, uzaktan "
        "calisma gunlerinde mesai saatleri icinde ulasilabilir olmalidir."
    )

    doc.add_heading("13.2 Ekipman ve Sorumluluklar", level=2)
    doc.add_paragraph(
        "Uzaktan calisan personele sirket tarafindan dizustu bilgisayar ve "
        "gerekli yazilim lisanslari saglanir. Internet baglantisi calisanin "
        "sorumluligundadir. Uzaktan calisma sirasinda sirket verilerinin "
        "guvenligi icin VPN kullanimi zorunludur. Calisan, sirket bilgilerini "
        "ucuncu kisilerle paylasmamakla yukumludur. Uzaktan calisma gunlerinde "
        "gunluk raporlama yapilmasi zorunludur."
    )

    # ══════════════════════════════════════════
    # BOLUM 14: IS GUVENLIGI
    # ══════════════════════════════════════════
    doc.add_heading("14. Is Guvenligi", level=1)

    doc.add_heading("14.1 Genel Kurallar", level=2)
    doc.add_paragraph(
        "6331 sayili Is Sagligi ve Guvenligi Kanunu geregince, tum calisanlar "
        "is guvenligi kurallarına uymakla yukumludur. Uretim alanlarinda "
        "kisisel koruyucu donanim (KKD) kullanimi zorunludur: baret, celik "
        "burunlu ayakkabi, koruyucu gozluk, eldiven ve gerekli hallerde "
        "kulak koruyucu. Is guvenligi kurallarina uymayan calisanlar hakkinda "
        "disiplin islemleri uygulanir."
    )

    doc.add_heading("14.2 Kaza Bildirimi", level=2)
    doc.add_paragraph(
        "Is kazasi veya ramak kala olayi meydana gelmesi halinde, derhal "
        "Is Guvenligi Uzmani ve IK departmanina bildirim yapilmalidir. "
        "Is kazasi bildirim formu 24 saat icinde doldurulmalidir. Agir is "
        "kazalari SGK'ya 3 is gunu icinde bildirilir. Her is kazasi sonrasinda "
        "kok neden analizi yapilir ve onleyici tedbirler belirlenir. Yilda "
        "en az 2 kez is guvenligi tatbikati (yangin, deprem) yapilir."
    )

    doc.add_heading("14.3 Saglik Kontrolleri", level=2)
    doc.add_paragraph(
        "Tum calisanlar icin yilda 1 kez periyodik saglik muayenesi yapilir. "
        "Tehlikeli ve cok tehlikeli sinifta calisan personel icin bu muayene "
        "6 ayda bir tekrarlanir. Gece vardiyasinda calisan personel icin "
        "ek saglik kontrolleri yapilabilir. Periyodik muayeneler is saatleri "
        "icinde yapilir ve bu sureler calisma suresinden sayilir. Muayene "
        "sonuclari gizlilik ilkesi kapsamında saklanir."
    )

    # ══════════════════════════════════════════
    # BOLUM 15: VERI GIZLILIGI / KVKK
    # ══════════════════════════════════════════
    doc.add_heading("15. Veri Gizliligi ve KVKK", level=1)

    doc.add_heading("15.1 Kisisel Veri Koruma", level=2)
    doc.add_paragraph(
        "6698 sayili Kisisel Verilerin Korunmasi Kanunu (KVKK) kapsaminda, "
        "Yorglass tum calisan ve aday verilerini kanuni sureler icinde ve "
        "belirtilen amaclarla isler. Calisan bilgileri (kimlik, iletisim, "
        "maas, saglik) yalnizca yetkili IK personeli tarafindan erisilebiir. "
        "Calisan verileri ucuncu kisilerle paylalilamaz (yasal zorunluluklar "
        "haric). Her calisan, KVKK kapsamindaki haklarini (bilgi edinme, "
        "duzeltme, silme) IK departmanina yazili basvuru ile kullanabilir."
    )

    doc.add_heading("15.2 Sirket Bilgi Guvenligi", level=2)
    doc.add_paragraph(
        "Calisanlar, gorevleri sirasinda edindikleri ticari sirlar, musteri "
        "bilgileri, uretim formulleri, fiyatlandirma stratejileri ve diger "
        "gizli bilgileri is sozlesmesi suresince ve sona erdikten sonra 2 yil "
        "boyunca gizli tutmakla yukumludur. Sirket bilgilerinin kisisel e-posta, "
        "sosyal medya veya diger kanallar araciligiyla paylailması yasaktir. "
        "Gizlilik ihlali durumunda yasal islem baslatilir."
    )

    # ══════════════════════════════════════════
    # BOLUM 16: YEMEK VE ULASIM YARDIMI
    # ══════════════════════════════════════════
    doc.add_heading("16. Yemek ve Ulasim Yardimi", level=1)

    doc.add_heading("16.1 Yemek Yardimi", level=2)
    doc.add_paragraph(
        "Yorglass, tum calisanlarina ogle yemegi yardimi saglar. Uretim "
        "tesislerinde sirket yemekhanesi mevcuttur ve yemekler ucretsiz "
        "olarak sunulur. Merkez ofis calisanlarina gunluk 150 TL tutarinda "
        "yemek karti (Sodexo/Multinet) verilir. Yemek karti tutari her yil "
        "Ocak ayinda guncellenir. Gece vardiyasinda calisan personele ek "
        "olarak gece yemegi saglanir."
    )

    doc.add_heading("16.2 Ulasim Yardimi", level=2)
    doc.add_paragraph(
        "Uretim tesislerinde calisan personel icin sirket servisi hizmeti "
        "sunulmaktadir. Servis guzergahlari her donem basindan personel "
        "yogunluguna göre belirlenir. Merkez ofis calisanlarina aylik 1.500 TL "
        "tutarinda ulasim yardimi yapilir. Saha gorevlerinde kullanilmak "
        "uzere sirket arac tahsisi yapilabilir; arac kullanim kurallari ayrica "
        "duzenlenmistir. Ulasim yardimi tutari her yil Ocak ayinda guncellenir."
    )

    # ══════════════════════════════════════════
    # BOLUM 17: IHBAR VE KIDEM TAZMINATI
    # ══════════════════════════════════════════
    doc.add_heading("17. Ihbar ve Kidem Tazminati", level=1)

    doc.add_heading("17.1 Ihbar Suresi", level=2)
    doc.add_paragraph(
        "4857 sayili Is Kanunu'nun 17. maddesi geregince, belirsiz sureli is "
        "sozlesmelerinin feshinden once bildirim (ihbar) surelerine uyulmasi "
        "zorunludur: 6 aya kadar kidem suresi icin 2 hafta, 6 aydan 1.5 yila "
        "kadar kidem suresi icin 4 hafta, 1.5 yildan 3 yila kadar kidem suresi "
        "icin 6 hafta, 3 yildan fazla kidem suresi icin 8 hafta. Ihbar suresine "
        "uymayan taraf, ihbar tazminati odemekle yukumludur."
    )

    doc.add_heading("17.2 Kidem Tazminati", level=2)
    doc.add_paragraph(
        "1475 sayili Is Kanunu'nun yururlukte olan 14. maddesi geregince, "
        "en az 1 yillik kidem suresini dolduran calisanlara kidem tazminati "
        "odenir. Kidem tazminati hesaplamasinda her tam yil icin 30 gunluk "
        "brut ucret esas alinir. Kidem tazminati tavanini asan kisim odenmez; "
        "tavan tutari her yil Hukumet tarafindan belirlenir. Istifa eden "
        "calisanlar kidem tazminatindan yararlanamazlar (emeklilik, askerlik, "
        "evlilik gibi ozel durumlar haric). Is sozlesmesinin isveren "
        "tarafindan hakli neden olmaksizin feshedilmesi halinde kidem "
        "tazminati odenmesi zorunludur."
    )

    doc.add_heading("17.3 Cikis Islemleri", level=2)
    doc.add_paragraph(
        "Isten ayrilma veya is akdi feshi durumunda calisan, zimmetindeki "
        "tum sirket mallarini (bilgisayar, telefon, arac, kimlik karti, anahtar) "
        "IK departmanina teslim etmekle yukumludur. Cikis islemleri IK "
        "departmani tarafindan koordine edilir ve iliskik kesme formu tum "
        "ilgili birimler tarafindan imzalanir. Kalan yillik izin ucretleri "
        "ve varsa fazla mesai alacaklari son aylik ile birlikte odenir. "
        "SGK cikis bildirgesi 10 gun icinde yapilir."
    )

    # Kaydet
    output_path = Path(__file__).resolve().parent / "data" / "yorglass_ik_prosedur.docx"
    doc.save(str(output_path))
    return str(output_path)


if __name__ == "__main__":
    path = create_ik_document()
    print(f"Dokuman olusturuldu: {path}")
